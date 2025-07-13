import io, os, datetime, math
import streamlit as st, pandas as pd, plotly.express as px
from docx import Document

# ── API key present? ────────────────────────────────────────────────
API = os.getenv("MP_API_KEY") or st.secrets.get("MP_API_KEY","")
if len(API)!=32:
    st.error("🛑  Set a valid 32-character MP_API_KEY")
    st.stop()

# ── backend helpers ─────────────────────────────────────────────────
from backend.perovskite_utils import (
    mix_abx3       as screen_binary,
    screen_ternary as screen_ternary,
    END_MEMBERS,
)

st.set_page_config("EnerMat Explorer v9.6", layout="wide")
st.title("🔬 EnerMat **Perovskite** Explorer v9.6")

# ── sidebar ────────────────────────────────────────────────────────
with st.sidebar:
    mode = st.radio("Choose screening type",["Binary A–B","Ternary A–B–C"])
    # end-members
    preset_A = st.selectbox("Preset A",END_MEMBERS,index=1)
    preset_B = st.selectbox("Preset B",END_MEMBERS,index=2)
    custom_A = st.text_input("Custom A (optional)")
    custom_B = st.text_input("Custom B (optional)")
    A = custom_A.strip() or preset_A
    B = custom_B.strip() or preset_B
    if mode=="Ternary A–B–C":
        preset_C = st.selectbox("Preset C",END_MEMBERS,index=3)
        custom_C = st.text_input("Custom C (optional)")
        C = custom_C.strip() or preset_C

    rh   = st.slider("Humidity [%]",0,100,50)
    temp = st.slider("Temperature [°C]",-20,100,25)
    bg_lo,bg_hi = st.slider("Gap window [eV]",0.5,3.0,(1.0,1.4),0.01)

    bow  = st.number_input("Bowing eV (neg⇒gap↑)",-1.0,1.0,-0.15,0.05)
    dx   = st.number_input("x-step",0.01,0.50,0.05,0.01)
    if mode=="Ternary A–B–C":
        dy = st.number_input("y-step",0.01,0.50,0.05,0.01)

# ── cached wrappers ─────────────────────────────────────────────────
@st.cache_data(show_spinner="⏳ binary …",max_entries=20)
def _run_binary(*args,**kw): return screen_binary(*args,**kw)

@st.cache_data(show_spinner="⏳ ternary …",max_entries=10)
def _run_ternary(*args,**kw): return screen_ternary(*args,**kw)

# ── run button ──────────────────────────────────────────────────────
if st.button("▶ Run screening",type="primary"):
    if mode=="Binary A–B":
        df = _run_binary(A,B,rh,temp,(bg_lo,bg_hi),bow,dx)
    else:
        df = _run_ternary(A,B,C,rh,temp,(bg_lo,bg_hi),
                          dict(AB=bow,AC=bow,BC=bow),dx,dy)
    if df.empty:
        st.error("No data – check formulas or API")
        st.stop()

    st.subheader("Results")
    st.dataframe(df,use_container_width=True,height=400)
    st.caption("Eox = oxidation energy (eV Sn⁻¹)")

    # simple scatter
    if "x" in df.columns and "Eox" in df.columns:
        pxfig = px.scatter(df,x="Eox",y="Eg",
                           color="Ehull",color_continuous_scale="Thermal",
                           hover_data=df.columns,width=1000,height=600)
        st.plotly_chart(pxfig,use_container_width=True)

    # downloads
    csv = df.to_csv(index=False).encode()
    st.download_button("📥 CSV",csv,"EnerMat_results.csv","text/csv")

    top = df.iloc[0]
    doc = Document(); doc.add_heading("EnerMat Report",0)
    for k,v in [("Date",datetime.date.today()),
                ("Top formula",top.formula),
                ("Eg (eV)",top.Eg),
                ("Ehull (eV/atom)",top.Ehull),
                ("Eox (eV Sn⁻¹)",top.Eox),
                ("Score",top.score)]:
        p=doc.add_paragraph(); p.add_run(f"{k}: ").bold=True; p.add_run(str(v))
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    st.download_button("📄 DOCX",bio,"EnerMat_report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Set parameters → Run screening")
