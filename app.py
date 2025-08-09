import datetime
import io
import json
from pathlib import Path
import base64

import streamlit as st
import pandas as pd
from plotly import graph_objects as go
import plotly.express as px
from docx import Document

from backend.perovskite_utils import (
    screen_binary,
    screen_ternary,
    END_MEMBERS,
    offline_mode,
)

# ─────────── STREAMLIT PAGE CONFIG ───────────
st.set_page_config("EnerMat Explorer – Lead-Free Perovskite PV Discovery Tool", layout="wide")
st.title("☀️ EnerMat Explorer | Lead-Free Perovskite PV Discovery Tool")

st.markdown(
    """
    <style>
      .css-1d391kg { border-right: 3px solid #0D47A1 !important; }
    </style>
    """,
    unsafe_allow_html=True,
)

# ─────────── SESSION STATE ───────────
if "history" not in st.session_state:
    st.session_state.history = []  # list of dicts: {"mode":..., "df":..., "params":...}

# ─────────── SIDEBAR ───────────
with st.sidebar:
    st.header("Mode")
    mode = st.radio("Choose screening type", ["Binary A–B", "Ternary A–B–C"])

    st.header("End-members")
    preset_A = st.selectbox("Preset A", END_MEMBERS, 0)
    preset_B = st.selectbox("Preset B", END_MEMBERS, 1)
    custom_A = st.text_input("Custom A (optional)").strip()
    custom_B = st.text_input("Custom B (optional)").strip()
    A = custom_A or preset_A
    B = custom_B or preset_B

    if mode.startswith("Ternary"):
        preset_C = st.selectbox("Preset C", END_MEMBERS, 2)
        custom_C = st.text_input("Custom C (optional)").strip()
        C = custom_C or preset_C

    st.header("Application")
    application = st.selectbox("Select application", ["single", "tandem", "indoor", "detector"])

    st.header("Environment (penalty hook)")
    rh = st.slider("Humidity [%]", 0, 100, 50)
    temp = st.slider("Temperature [°C]", -20, 100, 25)
    gamma_h = st.number_input("γ_h (RH weight, 0 = off)", 0.0, 2.0, 0.0, 0.05)
    gamma_t = st.number_input("γ_t (Temp weight, 0 = off)", 0.0, 2.0, 0.0, 0.05)

    st.header("Target band-gap [eV]")
    bg_lo, bg_hi = st.slider("Gap window", 0.50, 3.00, (1.00, 1.40), 0.01)

    st.header("Model settings")
    bow = st.number_input("Bowing (eV, negative ⇒ gap↑)", -1.0, 1.0, -0.15, 0.05)
    suggest_bow = st.checkbox("Suggest bowing from target (x=0.5)", value=False)
    dx = st.number_input("x-step", 0.01, 0.50, 0.05, 0.01)
    if mode.startswith("Ternary"):
        dy = st.number_input("y-step", 0.01, 0.50, 0.05, 0.01)

    z = st.slider(
        "Ge fraction z", 0.00, 0.80, 0.10, 0.05,
        help="B-site Ge²⁺ in CsSn₁₋zGe_zX₃"
    )

    st.header("Structural penalty")
    t0 = st.number_input("Target tolerance factor t₀", 0.80, 1.10, 0.95, 0.01)
    beta = st.number_input("Penalty stiffness β", 0.0, 5.0, 1.0, 0.1)

    # ── Clear history ──
    if st.button("🗑 Clear history"):
        st.session_state.history = []
        st.rerun()

    # ── Save/Load session ──
    st.subheader("Session")
    col_sav, col_ld = st.columns(2)
    with col_sav:
        if st.session_state.history:
            payload = json.dumps(st.session_state.history, default=lambda o: None)
            st.download_button("💾 Save session (.json)", payload, "EnerMat_session.json", "application/json")
    with col_ld:
        up = st.file_uploader("Load session", type=["json"], accept_multiple_files=False)
        if up:
            try:
                st.session_state.history = json.loads(up.read().decode("utf-8"))
                st.success("Session loaded.")
                st.rerun()
            except Exception as e:
                st.error(f"Failed to load session: {e}")

    # ── Developer credit in sidebar footer ──
    st.markdown(
        """
        <div style="font-size:0.85rem; color:#555; margin-top:0.5rem;">
          <strong>Developer:</strong> Dr Gbadebo Taofeek Yusuf (Academic World)<br>
          📞 +44 7776 727237 &nbsp; ✉️ das@academicworld.co.uk
        </div>
        """,
        unsafe_allow_html=True,
    )

# ─────────── WARNINGS / STATUS ───────────
if offline_mode:
    st.warning("Running in offline/demo mode (no valid MP API key found). Some formulas may return no data.")

# ─────────── CACHE WRAPPERS ───────────
@st.cache_data(show_spinner="⏳ Screening …", max_entries=20)
def _run_binary(args: dict):
    return screen_binary(**args)

@st.cache_data(show_spinner="⏳ Screening …", max_entries=10)
def _run_ternary(args: dict):
    return screen_ternary(**args)

# ─────────── OVERVIEW & HOW TO USE (unchanged styling) ───────────
st.markdown(
    """
    <style>
      .overview-box, .usage-box {
        background-color: #ffffff; border: 1px solid #dddddd; border-radius: 8px;
        padding: 24px; margin-bottom: 32px; color: #333333; font-family: Arial, sans-serif;
      }
      .overview-box h2, .usage-box h2 { margin-top: 0; color: #005FAD; }
    </style>
    <div class="overview-box">
      <h2>Context &amp; Scientific Justification</h2>
      <p>
        Lead–halide perovskites deliver record solar efficiencies but suffer from environmental toxicity and rapid degradation.
        Tin-based, lead-free analogues offer a safer path, yet optimising band gap (E<sub>g</sub>), phase stability (E<sub>hull</sub>),
        oxidation resistance (ΔE<sub>ox</sub>), and theoretical PCE remains challenging.
      </p>
      <p><em>EnerMat Explorer</em> combines calibrated gaps, hull energies, an oxidation proxy, and SQ-limit PCE into a sortable score.</p>
    </div>
    <div class="usage-box">
      <h2>🔧 How to Use EnerMat Explorer</h2>
      <ul>
        <li>Pick end-members A &amp; B (and C for ternary) or type custom formulas.</li>
        <li>Choose application, set band-gap window, bowing, steps, and optional Ge fraction.</li>
        <li>Optionally enable environmental/structural penalties and bowing suggestion.</li>
        <li>Run, compare, and export CSV/TXT/DOCX; save or load full sessions as JSON.</li>
      </ul>
    </div>
    """,
    unsafe_allow_html=True,
)

# ─────────── RUN / PREVIOUS ───────────
col_run, col_prev = st.columns([3, 1])
do_run  = col_run.button("▶ Run screening", type="primary")
do_prev = col_prev.button("⏪ Previous", disabled=not st.session_state.history)

if do_prev:
    st.session_state.history.pop()
    st.success("Showing previous result")
    if not st.session_state.history:
        st.stop()

if do_run:
    # NOTE: We no longer restrict A/B/C to END_MEMBERS; custom allowed.
    params = dict(
        rh=rh, temp=temp, application=application,
        gamma_h=gamma_h, gamma_t=gamma_t,
    )

    if mode.startswith("Binary"):
        args = dict(
            A=A, B=B, bg=(bg_lo, bg_hi), bow=bow, dx=dx, z=z,
            use_bowing_suggestion=bool(suggest_bow), t0=t0, beta=beta, **params
        )
        df = _run_binary(args)
    else:
        args = dict(
            A=A, B=B, C=C, bg=(bg_lo, bg_hi), bows={"AB": bow, "AC": bow, "BC": bow},
            dx=dx, dy=dy, z=z, **params
        )
        df = _run_ternary(args)

    if df.empty:
        st.error("No data returned for the chosen inputs. Try different end-members or widen the gap window.")
        st.stop()

    st.session_state.history.append({
        "mode": mode,
        "df": df,
        "params": {
            "A": A, "B": B, **({"C": C} if mode.startswith("Ternary") else {}),
            "application": application, "rh": rh, "temp": temp,
            "bg": (bg_lo, bg_hi), "bow": bow, "dx": dx, **({"dy": dy} if mode.startswith("Ternary") else {}),
            "z": z, "gamma_h": gamma_h, "gamma_t": gamma_t, "t0": t0, "beta": beta,
            "suggest_bow": suggest_bow
        }
    })

if not st.session_state.history:
    st.info("Press ▶ Run screening to begin.")
    st.stop()

# ─────────── DISPLAY RESULTS ───────────
df = st.session_state.history[-1]["df"]
mode = st.session_state.history[-1]["mode"]

tab_tbl, tab_plot, tab_dl = st.tabs(["📊 Table", "📈 Plot", "📥 Download"])

with tab_tbl:
    st.dataframe(df, use_container_width=True, height=460)

with tab_plot:
    if mode.startswith("Binary") and {"Ehull", "Eg", "score"}.issubset(df.columns):
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=df["Ehull"], y=df["Eg"], mode="markers",
            marker=dict(
                size=8 + 12 * df["score"], color=df["score"],
                colorscale="Viridis", cmin=0, cmax=1,
                colorbar=dict(title="Score"),
                line=dict(width=0.5, color="black"),
            ),
            hovertemplate="<b>%{customdata[6]}</b><br>"
                          "Eg=%{customdata[0]:.3f}±%{customdata[1]:.2f} eV<br>"
                          "Ehull=%{x:.4f}±%{customdata[2]:.2f} eV/at<br>"
                          "Score=%{marker.color:.3f} "
                          "[%{customdata[3]:.3f}, %{customdata[4]:.3f}]<br>"
                          "PCE_max=%{customdata[5]:.1f}%<extra></extra>",
            customdata=pd.DataFrame({
                "Eg": df["Eg"], "Eg_err": df.get("Eg_err", 0.0),
                "Ehull_err": df.get("Ehull_err", 0.0),
                "score_low": df.get("score_low", 0.0),
                "score_high": df.get("score_high", 0.0),
                "PCE": df.get("PCE_max (%)", 0.0),
                "label": df["formula"],
            }).to_numpy()
        ))
        fig.add_shape(
            type="rect", x0=0, x1=0.05, y0=bg_lo, y1=bg_hi,
            line=dict(color="LightSeaGreen", dash="dash"),
            fillcolor="LightSeaGreen", opacity=0.10
        )
        fig.update_layout(
            title="EnerMat Binary Screen",
            xaxis_title="Ehull (eV/atom)", yaxis_title="Eg (eV)",
            template="simple_white", width=720, height=540,
            margin=dict(l=60, r=60, t=60, b=60)
        )
        st.plotly_chart(fig, use_container_width=True)

    elif mode.startswith("Ternary") and {"x", "y", "score"}.issubset(df.columns):
        fig = px.scatter_3d(df, x="x", y="y", z="score", color="score",
                            color_continuous_scale="Viridis",
                            labels={"x": "B2 fraction", "y": "B3 fraction"},
                            height=520)
        st.plotly_chart(fig, use_container_width=True)

with tab_dl:
    # Current run CSV
    st.download_button("📥 Download current CSV", df.to_csv(index=False).encode(),
                       "EnerMat_results.csv", "text/csv")

    # Batch CSV of all runs
    if st.session_state.history:
        all_rows = []
        for i, h in enumerate(st.session_state.history, 1):
            dfi = pd.DataFrame(h["df"]).copy()
            for k, v in h.get("params", {}).items():
                dfi[f"param_{k}"] = v
            dfi["run_id"] = i
            all_rows.append(dfi)
        big = pd.concat(all_rows, ignore_index=True)
        st.download_button(
            "📦 Download ALL runs (CSV)",
            big.to_csv(index=False).encode(),
            "EnerMat_all_runs.csv", "text/csv"
        )

    # Auto TXT report (top row)
    _top = df.iloc[0]
    formula = str(_top["formula"])
    coords = ", ".join(
        f"{c}={_top[c]:.2f}"
        for c in ("x", "y", "z") if c in _top and pd.notna(_top[c])
    )
    label = formula if len(df) == 1 else f"{formula} ({coords})"

    _txt = (
        "EnerMat auto-report  "
        f"{datetime.date.today()}\n"
        f"Top candidate   : {label}\n"
        f"Band-gap [eV]   : {_top['Eg']} ± {_top.get('Eg_err', 0.0)}\n"
        f"Ehull [eV/at.]  : {_top['Ehull']} ± {_top.get('Ehull_err', 0.0)}\n"
        f"Eox_e [eV/e⁻]   : {_top.get('Eox_e', 'N/A')}\n"
        f"PCE_max [%]     : {_top.get('PCE_max (%)', 'N/A')}\n"
        f"Score           : {_top['score']} "
        f"[{_top.get('score_low', 'N/A')}, {_top.get('score_high', 'N/A')}]\n"
    )

    st.download_button("📄 Download TXT", _txt, "EnerMat_report.txt", mime="text/plain")

    # DOCX report
    _doc = Document()
    _doc.add_heading("EnerMat Report", level=0)
    _doc.add_paragraph(f"Date : {datetime.date.today()}")
    _doc.add_paragraph(f"Top candidate : {label}")

    table = _doc.add_table(rows=1, cols=2)
    table.style = "LightShading-Accent1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Property", "Value"
    for k in ("Eg", "Ehull", "Eox_e", "PCE_max (%)", "score", "score_low", "score_high"):
        if k in _top:
            row = table.add_row().cells
            row[0].text, row[1].text = k, str(_top[k])

    buf = io.BytesIO()
    _doc.save(buf)
    buf.seek(0)
    st.download_button("📝 Download DOCX", buf,
                       "EnerMat_report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
